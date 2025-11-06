# parser.py
import re
import io
from typing import Dict, List
import pdfplumber
import docx
import spacy

# load spaCy model at runtime (install and download en_core_web_sm separately)
nlp = spacy.load("en_core_web_sm")

EMAIL_RE = re.compile(r"[a-zA-Z0-9.+_-]+@[a-zA-Z0-9._-]+\.[a-zA-Z]+")
PHONE_RE = re.compile(r"(\+?\d{1,3}[\s-])?(?:\(?\d{2,4}\)?[\s-]?)?\d{3,4}[\s-]?\d{3,4}")

DEFAULT_SKILLS = {"python","java","c++","sql","postgresql","flask","django","aws","docker","kubernetes","spacy","nlp","react","javascript","node.js","html","css"}

def extract_text_from_pdf(path_or_bytes) -> str:
    text = []
    if isinstance(path_or_bytes, bytes):
        f = io.BytesIO(path_or_bytes)
        with pdfplumber.open(f) as pdf:
            for p in pdf.pages:
                text.append(p.extract_text() or "")
    else:
        with pdfplumber.open(path_or_bytes) as pdf:
            for p in pdf.pages:
                text.append(p.extract_text() or "")
    return "\n".join(text)
def extract_text_from_docx(path_or_bytes) -> str:
    if isinstance(path_or_bytes, bytes):
        f = io.BytesIO(path_or_bytes)
        doc = docx.Document(f)
    else:
        doc = docx.Document(path_or_bytes)
    return "\n".join([p.text for p in doc.paragraphs])


def normalize_whitespace(s: str) -> str:
    return re.sub(r"\s+\n", "\n", re.sub(r"[ \t]+", " ", s)).strip()


def find_emails(text: str) -> List[str]:
    return list(set(EMAIL_RE.findall(text)))


def find_phones(text: str) -> List[str]:
    phones = PHONE_RE.findall(text)
    if not phones:
        return []
    cleaned = set()
    for m in PHONE_RE.finditer(text):
        cleaned.add(m.group(0))
    return list(cleaned)


def extract_name(text: str) -> str:
    doc = nlp(text[:4000])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            candidate = ent.text.strip()
            if 1 <= len(candidate.split()) <= 4:
                return candidate
    first_line = text.strip().splitlines()[0] if text.strip().splitlines() else ""
    if len(first_line.split()) <= 6:
        return first_line.strip()
    return ""


def extract_skills(text: str, skills_set=None) -> List[str]:
    if skills_set is None:
        skills_set = DEFAULT_SKILLS
    text_low = text.lower()
    found = set()
    for skill in skills_set:
        if re.search(r"\b" + re.escape(skill.lower()) + r"\b", text_low):
            found.add(skill)
    return sorted(found)


def extract_education_sections(text: str) -> List[Dict]:
    ed_keywords = [
        "university",
        "college",
        "school",
        "bachelor",
        "master",
        "b.sc",
        "m.sc",
        "bachelors",
        "masters",
        "phd",
        "ph.d",
        "mba",
    ]
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    results = []
    for i, l in enumerate(lines):
        low = l.lower()
        if any(k in low for k in ed_keywords):
            raw = " ".join(lines[i : i + 3])
            parts = re.split(r"[-,–—]", raw)
            institution = parts[0].strip() if parts else ""
            degree = parts[1].strip() if len(parts) > 1 else ""
            results.append({"raw": raw, "institution": institution, "degree": degree})
    return results


def extract_experience_sections(text: str) -> List[Dict]:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    results = []
    for i, l in enumerate(lines):
        if " at " in l.lower() or re.search(r"\b\d{4}\b", l):
            raw = " ".join(lines[i : i + 4])
            if " at " in l.lower():
                parts = re.split(r"\bat\b", l, flags=re.IGNORECASE)
                title = parts[0].strip()
                company = parts[1].strip() if len(parts) > 1 else ""
            else:
                title = l
                company = ""
            results.append({"raw": raw, "title": title, "company": company})
    return results


def parse_resume(file_bytes: bytes, filename: str, skills_set=None) -> Dict:
    ext = filename.lower().split(".")[-1]
    if ext in {"pdf"}:
        text = extract_text_from_pdf(file_bytes)
    elif ext in {"docx", "doc"}:
        text = extract_text_from_docx(file_bytes)
    else:
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
    text = normalize_whitespace(text)
    emails = find_emails(text)
    phones = find_phones(text)
    name = extract_name(text)
    skills = extract_skills(text, skills_set)
    education = extract_education_sections(text)
    experience = extract_experience_sections(text)
    summary = text[:400].strip() if text else ""
    return {
        "full_text": text,
        "full_name": name,
        "emails": emails,
        "phones": phones,
        "skills": skills,
        "education": education,
        "experience": experience,
        "summary": summary,
    }


if __name__ == "__main__":
    with open("../sample_resume.pdf", "rb") as f:
        bytes_ = f.read()
    out = parse_resume(bytes_, "sample_resume.pdf")
    import json

    print(json.dumps(out, indent=2))